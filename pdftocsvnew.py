import os
import re
import csv
from docx import Document

# Directory containing Word documents
docx_folder = r"D:\Code\Invoices"
output_csv = r"D:\Code\InvoicesCSV\invoices_data.csv"

def extract_text_from_docx(docx_path):
    """Extracts text from a Word (.docx) document."""
    text = []
    doc = Document(docx_path)

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text.append(para.text.strip())

    return "\n".join(text), doc

def extract_invoice_details(text, doc):
    """Extracts invoice details using regex and retrieves PO Number from tables."""
    
    # Debugging Step 1: Print extracted text
    print(f"\n📝 Extracted Text:\n{text}")

    invoice_number_match = re.search(r"Invoice Number:\s*(.+)", text)
    due_date_match = re.search(r"Due Date:\s*(\d{2}/\d{2}/\d{4})", text)
    bill_to_match = re.search(r"Bill To:\s*(.+)", text)
    
    bill_to = bill_to_match.group(1).strip() if bill_to_match else "NOT FOUND"

    # Extract PO Number from tables
    po_number = "NOT FOUND"
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"📄 Table Row Data: {row_data}")  # Debugging Step 2
            
            po_number_cell = next((cell for cell in row_data if "PO Number" in cell), None)
            if po_number_cell:
                po_number = po_number_cell.split(":")[-1].strip()
                break

    # Extract Total Amount Due
    total_amount_match = re.search(r"Total Amount Due\s*\$([\d,]+\.\d{2})", text)
    total_amount = total_amount_match.group(1).strip() if total_amount_match else "NOT FOUND"

    extracted_data = {
        "Invoice Number": invoice_number_match.group(1).strip() if invoice_number_match else "NOT FOUND",
        "Due Date": due_date_match.group(1).strip() if due_date_match else "NOT FOUND",
        "Bill To": bill_to,
        "PO Number": po_number,
        "Total Amount Due": total_amount
    }

    print(f"📊 Extracted Data: {extracted_data}")  # Debugging Step 3
    return extracted_data

# Process all .docx files
print(f"📂 Files in directory: {os.listdir(docx_folder)}")  # Debugging Step 4
data_list = []

for filename in os.listdir(docx_folder):
    if filename.endswith(".docx"):
        docx_path = os.path.join(docx_folder, filename)
        try:
            print(f"\nProcessing: {filename}")  
            text, doc = extract_text_from_docx(docx_path)
            extracted_data = extract_invoice_details(text, doc)
            
            if any(value != "NOT FOUND" for value in extracted_data.values()):
                data_list.append(extracted_data)
            else:
                print(f"⚠️ No valid data found in {filename}")
        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")

# Ensure output directory exists
os.makedirs(os.path.dirname(output_csv), exist_ok=True)

# Debugging Step 5: Check final extracted data
print(f"\n📝 Final Data List: {data_list}")

# Write to CSV only if valid data exists
if data_list:
    with open(output_csv, mode="w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=["Invoice Number", "Due Date", "Bill To", "PO Number", "Total Amount Due"])
        writer.writeheader()
        writer.writerows(data_list)
    print(f"\n✅ Data successfully extracted and saved to {output_csv}")
else:
    print("\n⚠️ No valid data extracted. CSV file was not created.")